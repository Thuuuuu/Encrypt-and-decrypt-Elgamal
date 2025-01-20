import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from PIL import Image, ImageTk  
import random
import math
import time
from tkinter import filedialog
from docx import Document


def fast_exponentiation(a, x, p):
    kq = 1
    a %= p
    while x != 0:
        if x % 2 == 1:
            kq = (kq * a) % p
        a = (a * a) % p
        x //= 2
    return kq

def gcd(a, b):
    if b == 0:
        return a
    return gcd(b, a % b)

def chiver(n, a):
    if gcd(a, n) == 1:
        g0, g1 = n, a
        v0, v1 = 0, 1
        while g1 != 0:
            y = g0 // g1
            g0, g1 = g1, g0 - g1 * y
            v0, v1 = v1, v0 - v1 * y
        while v0 < 0:
            v0 += n
        return v0
    return -1

def encrypt_elgamal(p, a, y, m):
    value_m = ord(m)
    k = random.randint(1, p - 2)
    k_key = fast_exponentiation(y, k, p)
    c1 = fast_exponentiation(a, k, p)
    c2 = (k_key * value_m) % p
    return (c1, c2)

def decrypt_elgamal(c1, c2, x, p):
    k_key = fast_exponentiation(c1, x, p)
    k_inverse = chiver(p, k_key)
    result = (c2 * k_inverse) % p
    return chr(result)

p = random.randint(10000, 50000)
while not all(p % i != 0 for i in range(2, int(math.sqrt(p)) + 1)):
    p = random.randint(10000, 50000)

a = random.randint(2, p - 1)
while all(fast_exponentiation(a, (p - 1) // i, p) != 1 for i in range(2, int(math.sqrt(p - 1)) + 1)):
    a = random.randint(2, p - 1)

x = random.randint(1, p - 1)
y = fast_exponentiation(a, x, p)


def encrypt_message():
    plaintext = plaintext_entry.get()
    if not plaintext:
        messagebox.showerror("Error", "Please enter plaintext to encrypt.")
        return

    start_time = time.time() 
    ciphertext = [encrypt_elgamal(p, a, y, ch) for ch in plaintext]
    end_time = time.time()  

    ciphertext_entry.delete(0, tk.END)
    ciphertext_entry.insert(0, ' '.join([f"{c1},{c2}" for c1, c2 in ciphertext]))
    execution_time = (end_time - start_time ) *1000
    print( f"Thời gian mã hóa: {execution_time:.6f} ms")
    


def encrypt_file_doc():
    input_file = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")], title="Select a DOCX File to Encrypt")
    if not input_file:
        return

    output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], title="Save Encrypted DOCX File")
    if not output_file:
        return
    
    start_time = time.time()
    try:
        doc = Document(input_file)
        encrypted_doc = Document()

        for paragraph in doc.paragraphs:
            encrypted_paragraph = encrypted_doc.add_paragraph()
            for char in paragraph.text:
                c1, c2 = encrypt_elgamal(p, a, y, char)
                encrypted_paragraph.add_run(f"{c1} {c2} ")

        encrypted_doc.save(output_file)
        end_time = time.time() 
        execution_time = (end_time - start_time ) *1000
        print( f"Thời gian mã hóa file: {execution_time:.6f} ms")   
        messagebox.showinfo("Success", "Mã hóa file DOCX thành công!")
    except FileNotFoundError:
        messagebox.showerror("Error", "Không thể mở file đầu vào!")
    except Exception as e:
        messagebox.showerror("Error", f"Đã xảy ra lỗi: {e}")

def decrypt_file_doc():
    input_file = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")], title="Select a DOCX File to Encrypt")
    if not input_file:
        return

    output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], title="Save Encrypted DOCX File")
    if not output_file:
        return
    
    start_time = time.time()
    try:
        doc = Document(input_file)
        decrypted_doc = Document()

        for paragraph in doc.paragraphs:
            # Tạo một đoạn văn mới trong file DOCX giải mã
            decrypted_paragraph = decrypted_doc.add_paragraph()
            
            # Tách các cặp số từ đoạn văn bản
            elements = paragraph.text.split()
            i = 0
            while i < len(elements):
                try:
                    c1 = int(elements[i])
                    c2 = int(elements[i + 1])
                    decrypted_char = decrypt_elgamal(c1, c2, x, p)
                    decrypted_paragraph.add_run(decrypted_char)
                    i += 2
                except (ValueError, IndexError):
                    # Nếu không thể giải mã (lỗi cặp số), thêm ký tự "?" làm placeholder
                    decrypted_paragraph.add_run("?")
                    break
        
        # Lưu file DOCX đã giải mã
        decrypted_doc.save(output_file)
        end_time = time.time() 
        execution_time = (end_time - start_time)*10000
        print( f"Thời gian giải mã file: {execution_time:.6f} ms")
        messagebox.showinfo("Success", "Giải mã file DOCX thành công!")
    except FileNotFoundError:
        messagebox.showerror("Error", "Không thể mở file đầu vào!")
    except Exception as e:
        messagebox.showerror("Error", f"Đã xảy ra lỗi: {e}")
    

def generate_new_keys():
    global p, a, x, y

    # Tạo số nguyên tố p
    p = random.randint(10000, 50000)
    while not all(p % i != 0 for i in range(2, int(math.sqrt(p)) + 1)):
        p = random.randint(10000, 50000)

    # Tạo số nguyên g (a)
    a = random.randint(2, p - 1)
    while all(fast_exponentiation(a, (p - 1) // i, p) != 1 for i in range(2, int(math.sqrt(p - 1)) + 1)):
        a = random.randint(2, p - 1)

    # Tạo khóa bí mật x và khóa công khai y
    x = random.randint(1, p - 1)
    y = fast_exponentiation(a, x, p)

    # Cập nhật nhãn hiển thị khóa công khai và khóa bí mật
    public_key_label.config(text=f"Public Key: (y, a, p) = ({y}, {a}, {p})")
    private_key_label.config(text=f"Private Key: (x) = {x}")

    # Reset các trường nhập liệu
    reset_fields()
    messagebox.showinfo("Success", "Sinh khóa thành công!")


def decrypt_message():
    ciphertext = ciphertext_entry.get()
    if not ciphertext:
        messagebox.showerror("Error", "Please enter ciphertext to decrypt.")
        return
    try:
        start_time = time.time()  
        cipher_pairs = [tuple(map(int, pair.split(','))) for pair in ciphertext.split()]
        plaintext = ''.join(decrypt_elgamal(c1, c2, x, p) for c1, c2 in cipher_pairs)
        end_time = time.time()  

        execution_time = end_time - start_time  
        decryptext_entry.delete(0, tk.END)
        decryptext_entry.insert(0, plaintext)
        print(f"Thời gian giải mã: {execution_time:.6f} seconds")
    except ValueError:
        messagebox.showerror("Error", "Invalid ciphertext format.")

def save_current_data():
    plaintext = plaintext_entry.get()
    ciphertext = ciphertext_entry.get()
    decrypted_text = decryptext_entry.get()

    file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                             filetypes=[("Text files", "*.txt")],
                                             title="Save Current Data")
    if file_path:
        with open(file_path, "w") as file:
            if plaintext:
                file.write(f"Plaintext: {plaintext}\n")
            if ciphertext:
                file.write(f"Ciphertext: {ciphertext}\n")
            if decrypted_text:
                file.write(f"Decrypted Text: {decrypted_text}\n")
            file.write(f"Public Key (y, a, p): ({y}, {a}, {p})\n")
        messagebox.showinfo("Success", "Current data saved successfully.")

def load_data_from_file():
    file_path = filedialog.askopenfilename(defaultextension=".txt",
                                           filetypes=[("Text files", "*.txt")],
                                           title="Load Data File")
    if not file_path:
        return  
    
    try:
        with open(file_path, "r") as file:
            data = file.readlines()
        
        reset_fields()
        
        for line in data:
            if line.startswith("Plaintext:"):
                plaintext_entry.insert(0, line.split("Plaintext:")[1].strip())
            elif line.startswith("Ciphertext:"):
                ciphertext_entry.insert(0, line.split("Ciphertext:")[1].strip())
            elif line.startswith("Decrypted Text:"):
                decryptext_entry.insert(0, line.split("Decrypted Text:")[1].strip())
            elif line.startswith("Public Key (y, a, p):"):
                public_key_data = line.split("Public Key (y, a, p):")[1].strip()
                public_key_label.config(text=f"Public Key (y, a, p): {public_key_data}")
            elif line.startswith("Public Key (y, a, p):"):
                private_key_label = line.split("Public Key (x):")[1].strip()
                private_key_label.config(text=f"Public Key (x): {public_key_data}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load file: {e}")


def reset_fields():
    plaintext_entry.delete(0, tk.END)
    ciphertext_entry.delete(0, tk.END)
    decryptext_entry.delete(0, tk.END)


root = tk.Tk()
root.title("ElGamal Encryption/Decryption")


canvas = tk.Canvas(root, width=1200, height=600)
canvas.grid(row=0, column=0, columnspan=2)

save_button = tk.Button(root, text="Lưu File", command=save_current_data, font=("Helvetica", 14))
canvas.create_window(450, 400, window=save_button)

load_button = tk.Button(root, text="Đọc File", command=load_data_from_file, font=("Helvetica", 14))
canvas.create_window(750, 400, window=load_button)

load_button = tk.Button(root, text="Mã hóa file", command=encrypt_file_doc, font=("Helvetica", 14))
canvas.create_window(750, 500, window=load_button)

load_button = tk.Button(root, text="Giải Mã file", command=decrypt_file_doc, font=("Helvetica", 14))
canvas.create_window(450, 500, window=load_button)
      
public_key_label = tk.Label(root, text=f"Public Key: (y, a, p) = ({y}, {a}, {p})", font=("Helvetica", 14), bg="white")
canvas.create_window(600, 250, window=public_key_label)

private_key_label = tk.Label(root, text=f"Private Key: (x) = {x}", font=("Helvetica", 14), bg="white")
canvas.create_window(600, 300, window=private_key_label)

background_image = Image.open(r"D:\ki7\atbm\92805-Đinh Thị Thanh Thư-xây dựng chương trình mã hóa và giải mã Elgamal\anh nen chuong trinh.jpg")  # Đường dẫn tới ảnh của bạn
background_image = background_image.resize((1200, 600), Image.Resampling.LANCZOS)  # Sử dụng Resampling.LANCZOS
bg_photo = ImageTk.PhotoImage(background_image)

canvas.create_image(0, 0, image=bg_photo, anchor=tk.NW)
# canvas.lower()  


canvas.create_text(180, 70, text="Plaintext:", font=("Helvetica 14 bold"), anchor="w", fill="black")
plaintext_entry = tk.Entry(root, width=90,font=("Helvetica", 11))

canvas.create_text(180, 140, text="Ciphertext:", font=("Helvetica 14 bold"), anchor="w", fill="black")
ciphertext_entry = tk.Entry(root, width=90,font=("Helvetica", 11))


canvas.create_text(180, 200, text="Decryptext:", font=("Helvetica 14 bold"), anchor="w", fill="black")
decryptext_entry = tk.Entry(root, width=90, font=("Helvetica", 11))

plaintext_entry = tk.Entry(root, width=90, font=("Helvetica", 11))
canvas.create_window(700, 70, window=plaintext_entry)
ciphertext_entry = tk.Entry(root, width=90, font=("Helvetica", 11))
canvas.create_window(700, 140, window=ciphertext_entry)
decryptext_entry = tk.Entry(root, width=90, font=("Helvetica", 11))
canvas.create_window(700, 200, window=decryptext_entry)

encrypt_button = tk.Button(root, text="Mã hóa", command=encrypt_message,font=("Helvetica", 14))
canvas.create_window(500, 350, window=encrypt_button)

decrypt_button = tk.Button(root, text="Giải mã", command=decrypt_message,font=("Helvetica", 14))
canvas.create_window(700, 350, window=decrypt_button)

reset_button = tk.Button(root, text="Xóa", command=reset_fields,font=("Helvetica", 14))
canvas.create_window(600, 450, window=reset_button)

generate_key_button = tk.Button(root, text="Sinh Khóa Mới", command=generate_new_keys, font=("Helvetica", 14))
canvas.create_window(600, 500, window=generate_key_button)




root.mainloop()  